"""
Blueprint generation_contrats_bp.
Module de generation de contrats de travail a partir de modeles DOCX.

Onglet 1 : Formulaire de generation de contrat (remplacement des champs !XXXX!)
Onglet 2 : Gestion des modeles DOCX (upload, download, suppression)
"""
import os
import re
import unicodedata
from io import BytesIO
from flask import (Blueprint, render_template, request, redirect,
                   url_for, session, flash, send_file)
from database import get_db
from utils import login_required, get_setting, save_setting

generation_contrats_bp = Blueprint('generation_contrats_bp', __name__)

MODELES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'modeles_contrats'
)

CONTRATS_GENERES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'contrats_generes'
)

TYPES_CONTRAT = [
    ('CDI', 'CDI'),
    ('CDD_REMPLACEMENT', 'CDD de remplacement'),
    ('CDD_ACCROISSEMENT', "CDD accroissement d'activité"),
    ('CEE', 'CEE'),
    ('AUTRE', 'Autre'),
]

# Noms courts des critères ALISFA (pour les placeholders !CRITERE1! ... !CRITERE8!)
CRITERE_LABELS = [
    'Formation requise',
    "Complexité de l'emploi",
    'Autonomie',
    'Relationnel',
    'Responsabilité financière',
    'Responsabilité RH',
    'Sécurité',
    'Gestion de projet',
]

CRITERE_FIELDS = [
    'formation_niveau',
    'complexite_niveau',
    'autonomie_niveau',
    'relationnel_niveau',
    'finances_niveau',
    'rh_niveau',
    'securite_niveau',
    'projet_niveau',
]


def _get_modeles_dir():
    os.makedirs(MODELES_DIR, exist_ok=True)
    return MODELES_DIR


def _get_contrats_dir():
    os.makedirs(CONTRATS_GENERES_DIR, exist_ok=True)
    return CONTRATS_GENERES_DIR


def _peut_gerer():
    return session.get('profil') in ['comptable', 'directeur']


def _nettoyer_nom(texte):
    nfkd = unicodedata.normalize('NFKD', texte)
    sans_accents = ''.join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r'[^a-zA-Z0-9_-]', '', sans_accents)


def _chemin_securise(chemin, dossier):
    """Vérifie que le chemin est dans le dossier autorisé."""
    chemin_reel = os.path.realpath(chemin)
    dossier_reel = os.path.realpath(dossier)
    try:
        return os.path.commonpath([dossier_reel, chemin_reel]) == dossier_reel and chemin_reel != dossier_reel
    except ValueError:
        return False


@generation_contrats_bp.route('/generation_contrats')
@login_required
def generation_contrats():
    """Page principale avec deux onglets."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()

    # Données pour l'onglet 1 (formulaire)
    salaries = conn.execute('''
        SELECT u.id, u.nom, u.prenom, u.profil, u.adresse, u.date_naissance, u.numero_secu
        FROM users u
        WHERE u.actif = 1 AND u.profil != 'prestataire'
        ORDER BY u.nom, u.prenom
    ''').fetchall()

    postes = conn.execute(
        'SELECT id, intitule, total_points, formation_niveau, complexite_niveau, '
        'autonomie_niveau, relationnel_niveau, finances_niveau, rh_niveau, '
        'securite_niveau, projet_niveau FROM postes_alisfa ORDER BY intitule'
    ).fetchall()

    responsables = conn.execute('''
        SELECT id, nom, prenom, profil FROM users
        WHERE actif = 1 AND profil IN ('responsable', 'directeur')
        ORDER BY nom, prenom
    ''').fetchall()

    lieux = conn.execute('SELECT * FROM lieux_travail ORDER BY nom').fetchall()
    forfaits_cee = conn.execute('SELECT * FROM forfaits_cee ORDER BY montant').fetchall()
    modeles = conn.execute('SELECT * FROM modeles_contrats ORDER BY nom').fetchall()

    # Salaire socle (valeur par défaut 23000)
    salaire_socle = get_setting('salaire_socle') or '23000'

    # Onglet sélectionné
    onglet = request.args.get('onglet', '1')

    # Salarié pré-sélectionné (venant de la page infos_salaries)
    preselect_user_id = request.args.get('user_id', type=int)

    # Dernier contrat généré pour re-téléchargement
    dernier_contrat = None
    if preselect_user_id:
        dernier_contrat = conn.execute(
            'SELECT * FROM contrats_generes WHERE user_id = ? ORDER BY created_at DESC LIMIT 1',
            (preselect_user_id,)
        ).fetchone()

    conn.close()

    return render_template(
        'generation_contrats.html',
        salaries=salaries,
        postes=postes,
        responsables=responsables,
        lieux=lieux,
        forfaits_cee=forfaits_cee,
        modeles=modeles,
        salaire_socle=salaire_socle,
        types_contrat=TYPES_CONTRAT,
        critere_labels=CRITERE_LABELS,
        onglet=onglet,
        preselect_user_id=preselect_user_id,
        dernier_contrat=dernier_contrat,
    )


@generation_contrats_bp.route('/generation_contrats/generer', methods=['POST'])
@login_required
def generer_contrat():
    """Génère le contrat DOCX en remplaçant les placeholders."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    try:
        from docx import Document
    except ImportError:
        flash("La bibliothèque python-docx n'est pas installée.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats'))

    modele_id = request.form.get('modele_id', type=int)
    user_id = request.form.get('user_id', type=int)
    type_contrat = request.form.get('type_contrat', '').strip()

    if not modele_id or not user_id or not type_contrat:
        flash("Veuillez remplir tous les champs obligatoires.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    conn = get_db()

    # Récupérer le modèle
    modele = conn.execute('SELECT * FROM modeles_contrats WHERE id = ?', (modele_id,)).fetchone()
    if not modele:
        flash("Modèle introuvable.", 'error')
        conn.close()
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    # Récupérer le salarié
    salarie = conn.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    if not salarie:
        flash("Salarié introuvable.", 'error')
        conn.close()
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    # Récupérer le poste
    poste_id = request.form.get('poste_id', type=int)
    poste = None
    if poste_id:
        poste = conn.execute('SELECT * FROM postes_alisfa WHERE id = ?', (poste_id,)).fetchone()

    # Récupérer le responsable
    responsable_id = request.form.get('responsable_id', type=int)
    responsable = None
    if responsable_id:
        responsable = conn.execute('SELECT nom, prenom FROM users WHERE id = ?', (responsable_id,)).fetchone()

    # Récupérer les lieux (jusqu'à 3)
    lieux_ids = request.form.getlist('lieux_ids')
    lieux_noms = []
    for lid in lieux_ids[:3]:
        try:
            lid_int = int(lid)
        except (ValueError, TypeError):
            continue
        lieu = conn.execute('SELECT nom, adresse FROM lieux_travail WHERE id = ?', (lid_int,)).fetchone()
        if lieu:
            lieux_noms.append(lieu['nom'])

    # Salaire socle
    salaire_socle_val = request.form.get('salaire_socle', '').strip()
    if salaire_socle_val:
        save_setting('salaire_socle', salaire_socle_val)
    else:
        salaire_socle_val = get_setting('salaire_socle') or '23000'

    # Forfait CEE
    forfait_id = request.form.get('forfait_id', type=int)
    forfait_val = ''
    if forfait_id:
        forfait = conn.execute('SELECT montant, condition FROM forfaits_cee WHERE id = ?', (forfait_id,)).fetchone()
        if forfait:
            forfait_val = f"{forfait['montant']} €"

    conn.close()

    # Construire le dictionnaire de remplacement
    type_contrat_label = dict(TYPES_CONTRAT).get(type_contrat, type_contrat)
    is_cee = (type_contrat == 'CEE')
    is_cdd = type_contrat.startswith('CDD')
    is_cdd_remplacement = (type_contrat == 'CDD_REMPLACEMENT')

    # Pesée : récupérer les valeurs des critères du poste
    pesee_total = ''
    criteres_vals = {}
    if poste:
        pesee_total = str(poste['total_points']) if poste['total_points'] else ''
        for i, field in enumerate(CRITERE_FIELDS, 1):
            val = poste[field] if poste[field] is not None else ''
            criteres_vals[f'!CRITERE{i}!'] = str(val)

    # Horaires
    horaires = {}
    for jour in ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi']:
        horaires[f'!{jour.upper()}!'] = request.form.get(f'horaire_{jour}', '').strip()

    replacements = {
        '!NOM!': salarie['nom'] or '',
        '!PRENOM!': salarie['prenom'] or '',
        '!ADRESSE!': salarie['adresse'] or '',
        '!NAISSANCE!': _format_date(salarie['date_naissance']),
        '!SECURITESOCIALE!': salarie['numero_secu'] or '',
        '!TYPECONTRAT!': type_contrat_label,
        '!DEBUT!': _format_date(request.form.get('date_debut', '')),
        '!FIN!': _format_date(request.form.get('date_fin', '')) if is_cdd else '',
        '!REMPLACE!': request.form.get('remplace', '').strip() if is_cdd_remplacement else '',
        '!POSTE!': poste['intitule'] if poste else '',
        '!RESPONSABLE!': f"{responsable['nom']} {responsable['prenom']}" if responsable else '',
        '!HEBDO!': '' if is_cee else request.form.get('hebdo', '').strip(),
        '!ESSAI!': '' if is_cee else request.form.get('essai', '').strip(),
        '!LIEU!': ', '.join(lieux_noms) if lieux_noms else '',
        '!SOCLE!': '' if is_cee else salaire_socle_val,
        '!PESEE!': '' if is_cee else pesee_total,
        '!ANCIENNETE!': '' if is_cee else request.form.get('anciennete', '').strip(),
        '!FORFAIT!': forfait_val if is_cee else '',
        '!JOURS!': request.form.get('jours', '').strip() if is_cee else '',
    }
    replacements.update(criteres_vals)
    replacements.update(horaires)

    # Charger le modèle DOCX
    modele_path = os.path.join(_get_modeles_dir(), modele['fichier_path'])
    if not _chemin_securise(modele_path, _get_modeles_dir()) or not os.path.exists(modele_path):
        flash("Fichier modèle introuvable.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    try:
        doc = Document(modele_path)
        _remplacer_dans_document(doc, replacements)
    except Exception as e:
        flash(f"Erreur lors de la génération du contrat : {str(e)}", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1', user_id=user_id))

    # Construire le nom du fichier généré
    date_debut = request.form.get('date_debut', '')
    parts = date_debut.split('-') if date_debut else []
    date_fmt = f"{parts[2]}-{parts[1]}-{parts[0]}" if len(parts) == 3 else 'contrat'
    nom_clean = _nettoyer_nom(salarie['nom'])
    prenom_clean = _nettoyer_nom(salarie['prenom'])
    type_clean = _nettoyer_nom(type_contrat_label)
    nom_fichier = f"{date_fmt}_CONTRAT_{nom_clean}_{prenom_clean}_{type_clean}.docx"

    # Sauvegarder dans le dossier contrats_generes
    contrats_dir = _get_contrats_dir()
    chemin_sortie = os.path.join(contrats_dir, nom_fichier)

    # Supprimer l'ancien contrat généré pour ce salarié
    conn = get_db()
    ancien = conn.execute(
        'SELECT fichier_path FROM contrats_generes WHERE user_id = ? ORDER BY created_at DESC LIMIT 1',
        (user_id,)
    ).fetchone()
    if ancien and ancien['fichier_path']:
        ancien_chemin = os.path.join(contrats_dir, ancien['fichier_path'])
        if _chemin_securise(ancien_chemin, contrats_dir) and os.path.exists(ancien_chemin):
            os.remove(ancien_chemin)
    conn.execute('DELETE FROM contrats_generes WHERE user_id = ?', (user_id,))

    doc.save(chemin_sortie)

    conn.execute('''
        INSERT INTO contrats_generes (user_id, fichier_path, fichier_nom, type_contrat, created_by)
        VALUES (?, ?, ?, ?, ?)
    ''', (user_id, nom_fichier, nom_fichier, type_contrat_label, session['user_id']))
    conn.commit()
    conn.close()

    # Renvoyer le fichier directement
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=nom_fichier,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@generation_contrats_bp.route('/generation_contrats/retelecharger/<int:user_id>')
@login_required
def retelecharger_contrat(user_id):
    """Re-télécharge le dernier contrat généré pour un salarié."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()
    contrat = conn.execute(
        'SELECT * FROM contrats_generes WHERE user_id = ? ORDER BY created_at DESC LIMIT 1',
        (user_id,)
    ).fetchone()
    conn.close()

    if not contrat:
        flash("Aucun contrat généré pour ce salarié.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    contrats_dir = _get_contrats_dir()
    chemin = os.path.join(contrats_dir, contrat['fichier_path'])
    if not _chemin_securise(chemin, contrats_dir) or not os.path.exists(chemin):
        flash("Fichier introuvable.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    return send_file(chemin, as_attachment=True, download_name=contrat['fichier_nom'])


# ─── Gestion des modèles DOCX ───────────────────────────────────────────

@generation_contrats_bp.route('/generation_contrats/modele/upload', methods=['POST'])
@login_required
def upload_modele():
    """Uploader un modèle DOCX."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    nom = request.form.get('nom_modele', '').strip()
    fichier = request.files.get('fichier_modele')

    if not nom or not fichier or not fichier.filename:
        flash("Nom et fichier obligatoires.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    ext = os.path.splitext(fichier.filename)[1].lower()
    if ext != '.docx':
        flash("Le modèle doit être un fichier .docx.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    nom_clean = _nettoyer_nom(nom)
    nom_fichier = f"MODELE_{nom_clean}.docx"
    modeles_dir = _get_modeles_dir()
    chemin = os.path.join(modeles_dir, nom_fichier)

    # Gérer les doublons
    compteur = 1
    base = f"MODELE_{nom_clean}"
    while os.path.exists(chemin):
        nom_fichier = f"{base}_{compteur}.docx"
        chemin = os.path.join(modeles_dir, nom_fichier)
        compteur += 1

    fichier.save(chemin)

    conn = get_db()
    conn.execute(
        'INSERT INTO modeles_contrats (nom, fichier_path, fichier_nom, created_by) VALUES (?, ?, ?, ?)',
        (nom, nom_fichier, fichier.filename, session['user_id'])
    )
    conn.commit()
    conn.close()

    flash(f"Modèle '{nom}' ajouté.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))


@generation_contrats_bp.route('/generation_contrats/modele/remplacer/<int:modele_id>', methods=['POST'])
@login_required
def remplacer_modele(modele_id):
    """Remplacer un modèle DOCX existant."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    fichier = request.files.get('fichier_modele')
    if not fichier or not fichier.filename:
        flash("Veuillez sélectionner un fichier.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    ext = os.path.splitext(fichier.filename)[1].lower()
    if ext != '.docx':
        flash("Le modèle doit être un fichier .docx.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    conn = get_db()
    modele = conn.execute('SELECT * FROM modeles_contrats WHERE id = ?', (modele_id,)).fetchone()
    if not modele:
        flash("Modèle introuvable.", 'error')
        conn.close()
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    modeles_dir = _get_modeles_dir()
    ancien_chemin = os.path.join(modeles_dir, modele['fichier_path'])
    if _chemin_securise(ancien_chemin, modeles_dir) and os.path.exists(ancien_chemin):
        os.remove(ancien_chemin)

    fichier.save(ancien_chemin)
    conn.execute(
        'UPDATE modeles_contrats SET fichier_nom = ?, created_at = CURRENT_TIMESTAMP WHERE id = ?',
        (fichier.filename, modele_id)
    )
    conn.commit()
    conn.close()

    flash(f"Modèle '{modele['nom']}' remplacé.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))


@generation_contrats_bp.route('/generation_contrats/modele/telecharger/<int:modele_id>')
@login_required
def telecharger_modele(modele_id):
    """Télécharger un modèle DOCX."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()
    modele = conn.execute('SELECT * FROM modeles_contrats WHERE id = ?', (modele_id,)).fetchone()
    conn.close()

    if not modele:
        flash("Modèle introuvable.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    modeles_dir = _get_modeles_dir()
    chemin = os.path.join(modeles_dir, modele['fichier_path'])
    if not _chemin_securise(chemin, modeles_dir) or not os.path.exists(chemin):
        flash("Fichier introuvable sur le serveur.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    return send_file(chemin, as_attachment=True, download_name=modele['fichier_nom'])


@generation_contrats_bp.route('/generation_contrats/modele/supprimer/<int:modele_id>', methods=['POST'])
@login_required
def supprimer_modele(modele_id):
    """Supprimer un modèle DOCX."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()
    modele = conn.execute('SELECT * FROM modeles_contrats WHERE id = ?', (modele_id,)).fetchone()
    if not modele:
        flash("Modèle introuvable.", 'error')
        conn.close()
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))

    modeles_dir = _get_modeles_dir()
    chemin = os.path.join(modeles_dir, modele['fichier_path'])
    if _chemin_securise(chemin, modeles_dir) and os.path.exists(chemin):
        os.remove(chemin)

    conn.execute('DELETE FROM modeles_contrats WHERE id = ?', (modele_id,))
    conn.commit()
    conn.close()

    flash(f"Modèle '{modele['nom']}' supprimé.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='2'))


# ─── Gestion des lieux de travail ───────────────────────────────────────

@generation_contrats_bp.route('/generation_contrats/lieu/ajouter', methods=['POST'])
@login_required
def ajouter_lieu():
    """Ajouter un lieu de travail."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    nom = request.form.get('nom_lieu', '').strip()
    adresse = request.form.get('adresse_lieu', '').strip()

    if not nom or not adresse:
        flash("Nom et adresse obligatoires.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    conn = get_db()
    conn.execute(
        'INSERT INTO lieux_travail (nom, adresse, created_by) VALUES (?, ?, ?)',
        (nom, adresse, session['user_id'])
    )
    conn.commit()
    conn.close()

    flash(f"Lieu '{nom}' ajouté.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))


@generation_contrats_bp.route('/generation_contrats/lieu/supprimer/<int:lieu_id>', methods=['POST'])
@login_required
def supprimer_lieu(lieu_id):
    """Supprimer un lieu de travail."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()
    conn.execute('DELETE FROM lieux_travail WHERE id = ?', (lieu_id,))
    conn.commit()
    conn.close()

    flash("Lieu supprimé.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))


# ─── Gestion des forfaits CEE ────────────────────────────────────────────

@generation_contrats_bp.route('/generation_contrats/forfait_cee/ajouter', methods=['POST'])
@login_required
def ajouter_forfait_cee():
    """Ajouter un forfait CEE."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    try:
        montant = float(request.form.get('montant_forfait', '0').replace(',', '.'))
    except (ValueError, TypeError):
        flash("Montant invalide.", 'error')
        return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))

    condition = request.form.get('condition_forfait', '').strip() or None

    conn = get_db()
    conn.execute(
        'INSERT INTO forfaits_cee (montant, condition, created_by) VALUES (?, ?, ?)',
        (montant, condition, session['user_id'])
    )
    conn.commit()
    conn.close()

    flash(f"Forfait {montant} € ajouté.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))


@generation_contrats_bp.route('/generation_contrats/forfait_cee/supprimer/<int:forfait_id>', methods=['POST'])
@login_required
def supprimer_forfait_cee(forfait_id):
    """Supprimer un forfait CEE."""
    if not _peut_gerer():
        flash("Acces non autorise.", 'error')
        return redirect(url_for('dashboard_bp.dashboard'))

    conn = get_db()
    conn.execute('DELETE FROM forfaits_cee WHERE id = ?', (forfait_id,))
    conn.commit()
    conn.close()

    flash("Forfait supprimé.", 'success')
    return redirect(url_for('generation_contrats_bp.generation_contrats', onglet='1'))


# ─── Helpers ────────────────────────────────────────────────────────────

def _format_date(date_str):
    """Convertit YYYY-MM-DD en JJ/MM/AAAA."""
    if not date_str:
        return ''
    parts = date_str.strip().split('-')
    if len(parts) == 3:
        return f"{parts[2]}/{parts[1]}/{parts[0]}"
    return date_str


def _remplacer_dans_paragraphe(para, replacements):
    """
    Remplace les placeholders dans un paragraphe en gérant les runs fragmentés.
    Certains placeholders peuvent être répartis sur plusieurs runs consécutifs.
    Stratégie : assembler le texte de tous les runs, remplacer, puis mettre le
    résultat dans le premier run en conservant sa mise en forme.
    """
    if not para.runs:
        return

    # Assembler le texte complet du paragraphe
    texte_assemblé = ''.join(r.text for r in para.runs)
    nouveau_texte = texte_assemblé
    for placeholder, valeur in replacements.items():
        nouveau_texte = nouveau_texte.replace(placeholder, valeur)

    if nouveau_texte == texte_assemblé:
        # Rien à remplacer dans ce paragraphe
        return

    # Mettre le texte résultant dans le premier run, vider les suivants
    para.runs[0].text = nouveau_texte
    for run in para.runs[1:]:
        run.text = ''


def _remplacer_dans_document(doc, replacements):
    """Remplace tous les placeholders dans le document DOCX."""
    # Paragraphes du corps
    for para in doc.paragraphs:
        _remplacer_dans_paragraphe(para, replacements)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _remplacer_dans_paragraphe(para, replacements)

    # En-têtes et pieds de page
    for section in doc.sections:
        for para in section.header.paragraphs:
            _remplacer_dans_paragraphe(para, replacements)
        for para in section.footer.paragraphs:
            _remplacer_dans_paragraphe(para, replacements)
